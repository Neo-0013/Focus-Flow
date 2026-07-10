"""
DocForge Word Generator — word.py
Uses python-docx to produce real .docx files editable in Microsoft Word.
"""

from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from bs4 import BeautifulSoup, NavigableString
import datetime


# ── Theme colour maps for docx ───────────────────────────────────────
DOCX_THEMES = {
    "blue":   {"primary": (26, 35, 126),  "accent": (21, 101, 192), "light": (227, 242, 253)},
    "purple": {"primary": (74, 20, 140),  "accent": (123, 31, 162), "light": (243, 229, 245)},
    "green":  {"primary": (27, 94, 32),   "accent": (46, 125, 50),  "light": (232, 245, 233)},
    "red":    {"primary": (183, 28, 28),  "accent": (198, 40, 40),  "light": (255, 235, 238)},
    "dark":   {"primary": (26, 26, 46),   "accent": (22, 33, 62),   "light": (232, 234, 246)},
}


def _rgb(t: tuple) -> RGBColor:
    return RGBColor(*t)


def _cell_color(cell, rgb_tuple: tuple):
    """Apply background shading to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    hex_color = "{:02X}{:02X}{:02X}".format(*rgb_tuple)
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def generate_word(data: dict) -> bytes:
    """
    Generate a .docx document from the request data.
    Returns raw bytes to stream back to the browser.
    """
    theme_name = data.get("theme", "blue")
    theme = DOCX_THEMES.get(theme_name, DOCX_THEMES["blue"])
    font_size = data.get("font_size", 11)
    doc = Document()

    # ── Page setup ──────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Cm(21.0)   # A4
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.2)
    section.bottom_margin = Cm(2.2)

    # ── Custom styles ───────────────────────────────────────────
    _setup_styles(doc, theme, font_size)

    # ── Cover page ──────────────────────────────────────────────
    if data.get("show_cover", True):
        _build_cover(doc, data, theme, font_size)
        doc.add_page_break()

    # ── Metadata table ──────────────────────────────────────────
    _build_meta_table(doc, data, theme, font_size)

    # ── Main content ────────────────────────────────────────────
    html = data.get("html_content", "")
    if html and html.strip() not in ("", "<p></p>", "<p><br></p>"):
        _html_to_docx(doc, html, theme, font_size)
    else:
        p = doc.add_paragraph("Start writing your document in DocForge to see content here.")
        p.style = doc.styles["DF Body"]

    # ── Footer ──────────────────────────────────────────────────
    _add_footer(doc, data)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ══════════════════════════════════════════════════════════════
#  STYLE SETUP
# ══════════════════════════════════════════════════════════════

def _setup_styles(doc, theme, font_size):
    """Create custom named styles for consistent formatting."""
    styles = doc.styles

    def _add(name, base_name, size, bold=False, italic=False, color=None, align=None, space_before=0, space_after=6):
        try:
            st = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        except Exception:
            st = styles[name]
        st.base_style = styles[base_name]
        f = st.font
        f.name = "Calibri"
        f.size = Pt(size)
        f.bold = bold
        f.italic = italic
        if color:
            f.color.rgb = _rgb(color)
        pf = st.paragraph_format
        pf.space_before = Pt(space_before)
        pf.space_after  = Pt(space_after)
        if align:
            pf.alignment = align

    _add("DF Body",    "Normal", font_size,     color=(26, 26, 26), space_after=6)
    _add("DF H1",      "Normal", font_size + 7, bold=True, color=theme["primary"], space_before=14, space_after=6)
    _add("DF H2",      "Normal", font_size + 3, bold=True, color=theme["accent"],  space_before=12, space_after=4)
    _add("DF H3",      "Normal", font_size + 1, bold=True, italic=True, color=theme["accent"], space_before=8, space_after=3)
    _add("DF Bullet",  "Normal", font_size,     color=(26, 26, 26), space_after=3)
    _add("DF Ordered", "Normal", font_size,     color=(26, 26, 26), space_after=3)
    _add("DF Quote",   "Normal", font_size,     italic=True, color=(85, 85, 85), space_before=6, space_after=6)
    _add("DF Code",    "Normal", font_size - 1, color=(26, 35, 126))
    _add("DF Caption", "Normal", font_size - 2, italic=True, color=(100, 100, 100), align=WD_ALIGN_PARAGRAPH.CENTER)
    _add("DF Cover Title",  "Normal", 28, bold=True, color=(255, 255, 255), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    _add("DF Cover Sub",    "Normal", 12, color=(200, 225, 255), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    _add("DF Cover Meta",   "Normal", 9, color=(220, 235, 255), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)


# ══════════════════════════════════════════════════════════════
#  COVER PAGE
# ══════════════════════════════════════════════════════════════

def _build_cover(doc, data, theme, font_size):
    """Coloured cover block using table shading (closest docx equiv to Canvas drawing)."""
    # Coloured banner table spanning full width
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    _cell_color(cell, theme["primary"])

    # Institution
    inst = data.get("institution", "DocForge — FocusFlow").upper()
    p = cell.add_paragraph(inst)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0] if p.runs else p.add_run(inst)
    run.font.size = Pt(9)
    run.font.color.rgb = _rgb((180, 210, 255))
    run.font.name = "Calibri"

    # Spacer
    sp = cell.add_paragraph("")
    sp.add_run("\n" * 2)

    # Title
    title_p = cell.add_paragraph(data.get("title", "Untitled Document"))
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.runs[0] if title_p.runs else title_p.add_run(data.get("title", ""))
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    run.font.name = "Calibri"

    # Spacer
    cell.add_paragraph("\n")

    # Metadata lines
    meta_lines = []
    if data.get("author"):      meta_lines.append(f"Author: {data['author']}")
    if data.get("subject"):     meta_lines.append(f"Subject: {data['subject']}")
    date = data.get("date") or datetime.date.today().strftime("%B %d, %Y")
    meta_lines.append(f"Date: {date}")

    for line in meta_lines:
        mp = cell.add_paragraph(line)
        mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mrun = mp.runs[0] if mp.runs else mp.add_run(line)
        mrun.font.size = Pt(10)
        mrun.font.color.rgb = _rgb((220, 235, 255))
        mrun.font.name = "Calibri"

    cell.add_paragraph("\n\n")


# ══════════════════════════════════════════════════════════════
#  METADATA TABLE
# ══════════════════════════════════════════════════════════════

def _build_meta_table(doc, data, theme, font_size):
    rows_data = []
    if data.get("author"):      rows_data.append(("Author",      data["author"]))
    if data.get("subject"):     rows_data.append(("Subject",     data["subject"]))
    if data.get("institution"): rows_data.append(("Institution", data["institution"]))
    date = data.get("date") or datetime.date.today().strftime("%B %d, %Y")
    rows_data.append(("Date", date))

    if not rows_data:
        return

    tbl = doc.add_table(rows=len(rows_data), cols=2)
    tbl.style = "Table Grid"
    for i, (k, v) in enumerate(rows_data):
        key_cell = tbl.rows[i].cells[0]
        val_cell = tbl.rows[i].cells[1]

        # Alternating row colors
        if i % 2 == 1:
            _cell_color(key_cell, theme["light"])
            _cell_color(val_cell, theme["light"])

        kp = key_cell.paragraphs[0]
        kr = kp.add_run(k)
        kr.font.bold = True
        kr.font.size = Pt(font_size - 1)
        kr.font.color.rgb = _rgb((80, 80, 80))
        kr.font.name = "Calibri"

        vp = val_cell.paragraphs[0]
        vr = vp.add_run(v)
        vr.font.size = Pt(font_size - 1)
        vr.font.name = "Calibri"

    doc.add_paragraph("")  # spacer


# ══════════════════════════════════════════════════════════════
#  HTML → DOCX PARSER
# ══════════════════════════════════════════════════════════════

def _html_to_docx(doc, html: str, theme: dict, font_size: int):
    """Convert TipTap HTML to python-docx elements."""
    soup = BeautifulSoup(html, "html.parser")

    for el in soup.children:
        if isinstance(el, NavigableString):
            t = str(el).strip()
            if t:
                p = doc.add_paragraph(t)
                p.style = doc.styles["DF Body"]
            continue

        tag = el.name
        if not tag:
            continue

        if tag == "h1":
            _add_chapter_heading(doc, el.get_text(strip=True), theme, font_size)
        elif tag == "h2":
            _add_section_heading(doc, el.get_text(strip=True), theme, font_size)
        elif tag == "h3":
            p = doc.add_paragraph(el.get_text(strip=True))
            p.style = doc.styles["DF H3"]
        elif tag == "p":
            _add_rich_paragraph(doc, el, "DF Body")
        elif tag == "ul":
            for li in el.find_all("li", recursive=False):
                p = doc.add_paragraph(style="DF Bullet")
                p.add_run(f"•  {li.get_text(strip=True)}")
        elif tag == "ol":
            for i, li in enumerate(el.find_all("li", recursive=False), 1):
                p = doc.add_paragraph(style="DF Ordered")
                p.add_run(f"{i}.  {li.get_text(strip=True)}")
        elif tag == "blockquote":
            p = doc.add_paragraph(f'"{el.get_text(strip=True)}"')
            p.style = doc.styles["DF Quote"]
        elif tag == "pre":
            code_el = el.find("code")
            raw = (code_el or el).get_text()
            for line in raw.split("\n"):
                p = doc.add_paragraph(line or " ")
                p.style = doc.styles["DF Code"]
                p.runs[0].font.name = "Courier New"
        elif tag == "hr":
            doc.add_paragraph("─" * 60)
        elif tag == "table":
            _add_table_from_html(doc, el, theme, font_size)
        elif tag == "div":
            _html_to_docx(doc, str(el), theme, font_size)


def _add_chapter_heading(doc, text: str, theme: dict, font_size: int):
    """Coloured banner heading via table."""
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    _cell_color(cell, theme["primary"])
    p = cell.paragraphs[0]
    run = p.add_run(f"  {text}")
    run.font.bold = True
    run.font.size = Pt(font_size + 2)
    run.font.color.rgb = RGBColor(255, 255, 255)
    run.font.name = "Calibri"
    doc.add_paragraph("")


def _add_section_heading(doc, text: str, theme: dict, font_size: int):
    """Lighter sub-section heading."""
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    _cell_color(cell, theme["light"])
    p = cell.paragraphs[0]
    run = p.add_run(f"  {text}")
    run.font.bold = True
    run.font.size = Pt(font_size + 1)
    run.font.color.rgb = _rgb(theme["accent"])
    run.font.name = "Calibri"
    doc.add_paragraph("")


def _add_rich_paragraph(doc, el, style_name: str):
    """Add a paragraph preserving bold/italic/underline inline formatting."""
    p = doc.add_paragraph()
    try:
        p.style = doc.styles[style_name]
    except Exception:
        pass
    for child in el.children:
        if isinstance(child, NavigableString):
            run = p.add_run(str(child))
            run.font.name = "Calibri"
        else:
            t = child.get_text()
            run = p.add_run(t)
            run.font.name = "Calibri"
            if child.name in ("strong", "b"):
                run.bold = True
            if child.name in ("em", "i"):
                run.italic = True
            if child.name == "u":
                run.underline = True
            if child.name in ("s", "del"):
                run.font.strike = True
            if child.name == "code":
                run.font.name = "Courier New"


def _add_table_from_html(doc, table_el, theme, font_size):
    """Build a styled docx table from HTML table element."""
    headers = []
    rows = []
    thead = table_el.find("thead")
    tbody = table_el.find("tbody") or table_el
    if thead:
        headers = [th.get_text(strip=True) for th in thead.find_all("th")]
    for tr in tbody.find_all("tr"):
        row = [td.get_text(strip=True) for td in tr.find_all(["td", "th"])]
        if row:
            rows.append(row)
    if not headers and rows:
        headers = rows.pop(0)

    n_cols = max(len(headers), max((len(r) for r in rows), default=1), 1)
    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.style = "Table Grid"

    # Header row
    for j, h in enumerate(headers[:n_cols]):
        cell = tbl.rows[0].cells[j]
        _cell_color(cell, theme["accent"])
        run = cell.paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.size = Pt(font_size - 1)
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.name = "Calibri"

    # Data rows
    for i, row_data in enumerate(rows):
        bg = theme["light"] if i % 2 == 1 else (255, 255, 255)
        for j in range(n_cols):
            cell = tbl.rows[i + 1].cells[j]
            _cell_color(cell, bg)
            val = row_data[j] if j < len(row_data) else ""
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(font_size - 1)
            run.font.name = "Calibri"

    doc.add_paragraph("")


# ══════════════════════════════════════════════════════════════
#  FOOTER
# ══════════════════════════════════════════════════════════════

def _add_footer(doc, data):
    """Add simple footer with document title and page number."""
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        title_run = p.add_run(data.get("title", "")[:50] + "  |  ")
        title_run.font.size = Pt(8)
        title_run.font.color.rgb = RGBColor(130, 130, 130)
        title_run.font.name = "Calibri"

        # Page number field
        fldChar = OxmlElement("w:fldChar")
        fldChar.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.text = "PAGE"
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")

        run = p.add_run()
        run.font.size = Pt(8)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(130, 130, 130)
        run._r.append(fldChar)
        run._r.append(instrText)
        run._r.append(fldChar2)
